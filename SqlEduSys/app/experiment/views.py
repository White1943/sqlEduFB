from datetime import datetime

from app.LLM.model import get_llm_response
from app.LLM.prompts import generate_prompt
from flask import request, jsonify, current_app, Response, send_file
from . import experiment_bp
from app.models.experiments import Experiment
from app import db
from app.utils.response import ApiResponse
from sqlalchemy import or_
import json
from app.models.nl_queries import NLQueries
from app.models.schema import Sqls
from app.models.knowledges import KnowledgePoint
from app.LLM.services import get_experiment_report_response
import sys
sys.setrecursionlimit(100000)
import random
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import html2text
import re
import base64
from flask_cors import cross_origin
from bs4 import BeautifulSoup

@experiment_bp.route('/generate', methods=['POST'])
def generate_experiment():
    """生成实验报告"""
    try:
        data = request.get_json()
        schema_ids = data.get('schema_ids', [])
        experimentCount = data.get('count', 1)  # 获取生成数量
        points = data.get('points', [])
        version = data.get('version', 'student')  # 获取版本信息
        print(points)
        # print(count)
        # print(schema_ids)
        print(data)

        if not schema_ids or not points:
            return ApiResponse.error(message="请选择数据库模式和知识点")
        schemas = Sqls.query.filter(Sqls.id.in_(schema_ids)).all()
        if not schemas:
            return ApiResponse.error(message="未找到选中的数据库模式")

        schema_info = "\n".join([schema.file_content for schema in schemas])  
        generated_queries = []
        nl_queries = []
        for point in points:
            point_id = point.get('id')
            generateCount = point.get('generateCount', 1)
            
            knowledge_points = KnowledgePoint.query.get(point_id)
            if not knowledge_points: 
                print("知识点不存在")
                continue
            print(knowledge_points.point_name,point_id)

 
            nl_query= NLQueries.query.filter(
                NLQueries.knowledge_point_id == knowledge_points.id,
                NLQueries.status == 'approved',
                db.or_(*[db.text(f"FIND_IN_SET('{schema_id}', schema_ids) > 0") for schema_id in schema_ids])
            ).all()   
            nl_queries.extend(nl_query)
        
        dataExperiment   = ""
        for query in nl_queries:
            dataExperiment += f"<strong>题目:</strong> {query.query_text}<br><strong>答案:</strong> {query.generated_sql}<br><br>"
        print(dataExperiment)
        schemas = Sqls.query.filter(Sqls.id.in_(schema_ids)).all()
        # print("schemas ",schemas)

        schema_info = "\n".join([schema.file_content for schema in schemas])
        print( "schema_info" ,schema_info)
        print("schema_info is ok")
        # 下面的代码执行有问题
        for i in range(experimentCount):

            print("轮次",i)

        # 不是为每个知识点生成一个实验报告，而是一个实验报告就得包含所有前端传入的知识点，相应数量为generateCount



            teacher_prompt = generate_prompt(schema_info, points, 'teacher', dataExperiment)
            teacher_response = get_experiment_report_response(teacher_prompt)

            if teacher_response is None:
                return ApiResponse.error(message="教师版报告生成失败")

            current_app.logger.info(f"教师版报告内容长度: {len(teacher_response)}")
            # current_app.logger.info(f"准备插入的内容: {schema_ids}, {points}")
            # current_app.logger.info(f"教师版报告内容: {teacher_response.strip()}")
            # print(datetime.now().strftime('%Y-%m-%d-%H-%M-%S'))
            # 创建教师版实验记录
            # print(teacher_response.strip())
            print("打印教师实验内容")
            try:
                teacher_experiment = Experiment(
                    title=f"教师版实验报告 - {datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}",
                    description="根据选中的数据库模式生成的教师版实验报告",
                    notes=json.dumps({"schema_ids": schema_ids, "points": points}),
                    report_content=teacher_response.strip(),
                    version='teacher',
                    points_category_id=None
                )
                db.session.add(teacher_experiment)
                db.session.commit()
                current_app.logger.info("教师版实验记录插入成功")
            except Exception as e:
                current_app.logger.error(f"插入教师版实验记录失败: {str(e)}")
                return ApiResponse.error(message="插入教师版实验记录失败")

            # 生成学生版报告
            student_prompt = generate_prompt(schema_info, points, 'student', dataExperiment)
            student_response = get_experiment_report_response(student_prompt)

            if student_response is None:

                return ApiResponse.error(message="学生版报告生成失败")


            current_app.logger.info(f"学生版报告内容长度: {len(student_response)}")

            try:
                student_experiment = Experiment(
                    title=f"学生版实验报告 -   ",
                    description="根据选中的数据库模式生成的学生版实验报告",
                    content=json.dumps({"schema_ids": schema_ids, "points": points}),
                    report_content=student_response.strip(),
                    version='student',
                    points_category_id=None
                )
                db.session.add(student_experiment)
                db.session.commit()
                current_app.logger.info("学生版实验记录插入成功")
            except Exception as e:
                current_app.logger.error(f"插入学生版实验记录失败: {str(e)}")
                return ApiResponse.error(message="插入学生版实验记录失败")
         
        return ApiResponse.success(message="实验报告生成成功", data={"message": "所有实验报告已成功生成"})
    except Exception as e:
        current_app.logger.error(f"处理请求时发生错误: {str(e)}")
        return ApiResponse.error(message="处理请求时发生错误")

 

@experiment_bp.route('/download/<int:experiment_id>', methods=['GET'])
@cross_origin()
def download_experiment(experiment_id):
    """下载实验报告"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        
        experiment = Experiment.query.get_or_404(experiment_id)
        current_app.logger.info(f"找到实验报告: {experiment.title}")
        
        # 创建一个新的文档
        doc = Document()
        
        # 设置文档样式
        styles = doc.styles
        style_heading1 = styles['Heading 1']
        style_heading1.font.size = Pt(18)
        style_heading1.font.bold = True
        
        style_heading2 = styles['Heading 2']
        style_heading2.font.size = Pt(16)
        style_heading2.font.bold = True
        
        style_heading3 = styles['Heading 3']
        style_heading3.font.size = Pt(14)
        style_heading3.font.bold = True
        
        # 添加标题
        title = doc.add_heading(experiment.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加描述
        if experiment.description:
            desc_para = doc.add_paragraph()
            desc_para.add_run(experiment.description).italic = True
            doc.add_paragraph() # 空行
        
        # 解析HTML内容
        html_content = experiment.report_content
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 处理内容
        for element in soup.recursiveChildGenerator():
            if element.name == 'h1':
                doc.add_heading(element.text.strip(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.text.strip(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.text.strip(), level=3)
            elif element.name == 'h4':
                doc.add_heading(element.text.strip(), level=4)
            elif element.name == 'p':
                # 检查段落是否只包含一个图片
                images = element.find_all('img')
                if len(images) == 1 and not element.text.strip():
                    # 处理图片
                    img = images[0]
                    src = img.get('src', '')
                    
                    if src.startswith('data:image'):
                        # 处理Base64编码的图片
                        try:
                            # 解析图片类型和数据
                            content_type, image_data = src.split(';base64,')
                            image_type = content_type.split(':')[1]
                            
                            # 解码Base64数据
                            import base64
                            from io import BytesIO
                            
                            img_data = base64.b64decode(image_data)
                            img_stream = BytesIO(img_data)
                            
                            # 添加图片到文档
                            doc.add_picture(img_stream, width=Inches(6))
                            img_para = doc.paragraphs[-1]
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as img_err:
                            current_app.logger.error(f"处理图片失败: {str(img_err)}")
                            doc.add_paragraph(f"[图片处理失败]")
                else:
                    # 正常段落
                    para = doc.add_paragraph()
                    text = element.text.strip()
                    if text:  # 只有当有文本时才添加
                        para.add_run(text)
            elif element.name == 'ul' or element.name == 'ol':
                # 处理列表
                for li in element.find_all('li', recursive=False):
                    para = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                    para.add_run(li.text.strip())
            elif element.name == 'pre':
                # 处理代码块
                code = element.text.strip()
                if code:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run(code)
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(10)
                    code_para_format = code_para.paragraph_format
                    code_para_format.left_indent = Pt(36)
                    code_para_format.space_before = Pt(12)
                    code_para_format.space_after = Pt(12)

        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # 发送文件
        response = send_file(
            file_stream,
            as_attachment=True,
            download_name=f"实验报告_{experiment_id}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
    except Exception as e:
        current_app.logger.error(f"下载实验报告失败: {str(e)}")
        import traceback
        current_app.logger.error(traceback.format_exc())
        return ApiResponse.error(message=f"下载失败: {str(e)}")

@experiment_bp.route('/points/category/<int:category_id>', methods=['GET'])
def get_experiment_by_points_category(category_id):
    """获取实验报告"""
    experiments = Experiment.query.filter(Experiment.points_category_id == category_id).all()
    return ApiResponse.success(data=experiments)

@experiment_bp.route('/queries', methods=['GET'])
def get_experiments():
    try:
        # 获取查询参数
        schema_ids = request.args.get('schema_ids', '').split(',')
        page = request.args.get('page', 1, type=int)
        page_size = request.args.get('page_size', 10, type=int)
        
        if not schema_ids or schema_ids[0] == '':
            return ApiResponse.success(data={'items': [], 'total': 0})
            
        # 构建查询条件
        conditions = []
        for schema_id in schema_ids:
            conditions.append(db.text(f"FIND_IN_SET('{schema_id}', schema_ids) > 0"))
            
        # 构建查询
        query = Experiment.query.filter(db.or_(*conditions))
        
        # 获取总数
        total = query.count()
        
        # 分页
        paginated = query.order_by(Experiment.id.desc()).paginate(
            page=page,
            per_page=page_size,
            error_out=False
        )
        
        return ApiResponse.success(data={
            'items': [item.to_dict() for item in paginated.items],
            'total': total
        })
        
    except Exception as e:
        current_app.logger.error(f"获取查询列表失败: {str(e)}")
        return ApiResponse.error(message=str(e))

@experiment_bp.route('/generate2nd', methods=['POST'])
def generate_experiment2nd():
    """使用后端随机算法生成实验报告（不依赖大模型）"""
    try:
        data = request.get_json()
        schema_ids = data.get('schema_ids', [])
        experimentCount = data.get('count', 1)  # 获取生成数量
        points = data.get('points', [])

        if not schema_ids or not points:
            return ApiResponse.error(message="请选择数据库模式和知识点")
        
        # 获取选中的数据库模式
        schemas = Sqls.query.filter(Sqls.id.in_(schema_ids)).all()
        if not schemas:
            return ApiResponse.error(message="未找到选中的数据库模式")

        # 生成指定数量的实验报告
        for i in range(experimentCount):
            # 生成教师版报告
            teacher_content = generate_markdown_experiment(schemas, points, 'teacher')
            if not teacher_content:
                return ApiResponse.error(message="教师版报告生成失败")

            try:
                teacher_experiment = Experiment(
                    title=f"教师版实验报告(自动生成) - {datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}",
                    description="根据选中的数据库模式生成的教师版实验报告",
                    notes=json.dumps({"schema_ids": schema_ids, "points": points}),
                    report_content=teacher_content,
                    version='teacher',
                    points_category_id=None
                )
                db.session.add(teacher_experiment)
                db.session.commit()
                current_app.logger.info("教师版实验记录插入成功")
            except Exception as e:
                current_app.logger.error(f"插入教师版实验记录失败: {str(e)}")
                return ApiResponse.error(message="插入教师版实验记录失败")

            # 生成学生版报告
            student_content = generate_markdown_experiment(schemas, points, 'student')
            if not student_content:
                return ApiResponse.error(message="学生版报告生成失败")

            try:
                student_experiment = Experiment(
                    title=f"学生版实验报告(自动生成) - {datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}",
                    description="根据选中的数据库模式生成的学生版实验报告",
                    notes=json.dumps({"schema_ids": schema_ids, "points": points}),
                    report_content=student_content,
                    version='student',
                    points_category_id=None
                )
                db.session.add(student_experiment)
                db.session.commit()
                current_app.logger.info("学生版实验记录插入成功")
            except Exception as e:
                current_app.logger.error(f"插入学生版实验记录失败: {str(e)}")
                return ApiResponse.error(message="插入学生版实验记录失败")

        return ApiResponse.success(message="实验报告生成成功", data={"message": "所有实验报告已成功生成"})
    except Exception as e:
        current_app.logger.error(f"处理请求时发生错误: {str(e)}")
        return ApiResponse.error(message="处理请求时发生错误")

def generate_markdown_experiment(schemas, points, version):
    """
    生成markdown格式的实验报告
    
    Args:
        schemas: 数据库模式列表
        points: 知识点列表，包含id和数量
        version: 版本 ('teacher' 或 'student')
        
    Returns:
        str: markdown格式的实验报告内容
    """
    try:
        # 1. 生成报告头部
        now = datetime.now().strftime('%Y-%m-%d')
        content = []
        
        # 标题
        content.append(f"# SQL实验报告 {'(教师版)' if version == 'teacher' else '(学生版)'}")
        content.append(f"## 实验日期: {now}")
        content.append("")
        
        # 2. 数据库表结构部分
        content.append("## 一、数据库表结构")
        content.append("")
        
        for schema in schemas:
            content.append(f"### {schema.filename}")
            content.append("")
            content.append("```sql")
            content.append(schema.file_content)
            content.append("```")
            content.append("")
        
        # 3. 实验目的与要求
        content.append("## 二、实验目的与要求")
        content.append("")
        content.append("本次实验旨在帮助学生掌握以下SQL知识点和技能：")
        content.append("")
        
        for point_data in points:
            point_id = point_data.get('id')
            point = KnowledgePoint.query.get(point_id)
            if point:
                content.append(f"- **{point.point_name}**: {point.description}")
        
        content.append("")
        content.append("通过本次实验，学生将能够理解并应用上述SQL概念，提高数据库查询和操作能力。")
        content.append("")
        
        # 4. 实验内容与步骤
        content.append("## 三、实验内容与步骤")
        content.append("")
        
        question_number = 1
        
        # 为每个知识点获取相应的查询
        for point_data in points:
            point_id = point_data.get('id')
            count = point_data.get('generateCount', 1)
            
            point = KnowledgePoint.query.get(point_id)
            if not point:
                continue
                
            # 获取与该知识点相关的已批准的查询
            queries = NLQueries.query.filter(
                NLQueries.knowledge_point_id == point_id,
                NLQueries.status == 'approved',
                NLQueries.generated_sql.isnot(None),  # 确保有SQL语句
                db.or_(*[db.text(f"FIND_IN_SET('{schema.id}', schema_ids) > 0") for schema in schemas])
            ).all()
            
            # 如果查询数量不足，则返回所有可用查询
            selected_queries = random.sample(queries, min(count, len(queries))) if queries else []
            
            if selected_queries:
                content.append(f"### {point.point_name}")
                content.append("")
                content.append(f"{point.description}")
                content.append("")
                
                for query in selected_queries:
                    content.append(f"#### 问题 {question_number}: {query.query_text}")
                    content.append("")
                    
                    if version == 'teacher':
                        content.append("**答案：**")
                        content.append("")
                        content.append("```sql")
                        content.append(query.generated_sql)
                        content.append("```")
                        content.append("")
                        content.append("**解析：**")
                        content.append("")
                        content.append(f"本题考察了 {point.point_name} 的应用。" +
                                     f"SQL 语句使用了 {point.description}。" +
                                     "通过该查询，可以获取符合条件的数据。")
                        content.append("")
                    else:
                        content.append("请在下方空白处编写SQL语句：")
                        content.append("")
                        content.append("```sql")
                        content.append("-- 在此处编写答案")
                        content.append("```")
                        content.append("")
                    
                    question_number += 1
        
        # 5. 总结
        content.append("## 四、实验总结")
        content.append("")
        
        if version == 'teacher':
            content.append("本次实验主要涵盖了以下知识点：")
        else:
            content.append("通过本次实验，请总结你学到的SQL知识点和技能：")
            
        content.append("")
        for point_data in points:
            point_id = point_data.get('id')
            point = KnowledgePoint.query.get(point_id)
            if point:
                content.append(f"- **{point.point_name}**")
        
        if version == 'teacher':
            content.append("")
            content.append("希望同学们通过本次实验，能够更加熟练地掌握SQL查询的基本语法和常用技巧，为进一步学习复杂的SQL查询打下坚实的基础。")
        else:
            content.append("")
            content.append("_（请在此处填写你的实验总结和收获）_")
        
        return "\n".join(content)
    except Exception as e:
        current_app.logger.error(f"生成markdown实验报告失败: {str(e)}")
        return None

# 获取实验报告列表API
@experiment_bp.route('/list', methods=['GET'])
def get_experiments_list():
    try:
        # 获取查询参数
        page = request.args.get('page', 1, type=int)
        page_size = request.args.get('page_size', 10, type=int)
        version = request.args.get('version')
        
        # 构建查询
        query = Experiment.query
        
        # 按版本筛选
        if version in ['teacher', 'student']:
            query = query.filter(Experiment.version == version)
        
        # 获取总数
        total = query.count()
        
        # 分页
        paginated = query.order_by(Experiment.id.desc()).paginate(
            page=page,
            per_page=page_size,
            error_out=False
        )
        
        return ApiResponse.success(data={
            'items': [item.to_dict() for item in paginated.items],
            'total': total
        })
        
    except Exception as e:
        current_app.logger.error(f"获取实验报告列表失败: {str(e)}")
        return ApiResponse.error(message=str(e))

@experiment_bp.route('/update/<int:experiment_id>', methods=['PUT'])
def update_experiment(experiment_id):
    """更新实验报告"""
    try:
        experiment = Experiment.query.get_or_404(experiment_id)
        data = request.get_json()
        
        # 更新可编辑字段
        if 'title' in data:
            experiment.title = data['title']
        if 'description' in data:
            experiment.description = data['description']
        if 'report_content' in data:
            # 存储HTML内容
            experiment.report_content = data['report_content']
        if 'objectives' in data:
            experiment.objectives = data['objectives']
        if 'requirements' in data:
            experiment.requirements = data['requirements']
        if 'steps' in data:
            experiment.steps = data['steps']
        if 'notes' in data:
            experiment.notes = data['notes']
        
        db.session.commit()
        return ApiResponse.success(message="实验报告更新成功")
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"更新实验报告失败: {str(e)}")
        return ApiResponse.error(message=f"更新失败: {str(e)}")

@experiment_bp.route('/download_html/<int:experiment_id>', methods=['GET'])
def download_experiment_html(experiment_id):
    """下载HTML格式的实验报告"""
    try:
        experiment = Experiment.query.get_or_404(experiment_id)
        
        # 获取报告内容
        content = experiment.report_content
        
        # 如果内容是Markdown格式，转换为HTML
        if not ('<' in content and '>' in content):
            content = markdown.markdown(content)
        
        # 创建完整的HTML文档
        html = f"""<!DOCTYPE html>
<html>
    <head>
        <meta charset="utf-8">
        <title>{experiment.title}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            h1 {{ color: #333; text-align: center; }}
            pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 5px; overflow-x: auto; }}
            code {{ font-family: Consolas, Monaco, monospace; }}
            img {{ max-width: 100%; height: auto; }}
        </style>
    </head>
    <body>
        <h1>{experiment.title}</h1>
        {content}
    </body>
</html>"""
        
        # 使用send_file发送字符串流
        html_stream = io.BytesIO(html.encode('utf-8'))
        
        return send_file(
            html_stream,
            as_attachment=True,
            download_name=f"实验报告_{experiment_id}.html",
            mimetype='text/html'
        )
    except Exception as e:
        current_app.logger.error(f"下载HTML报告失败: {str(e)}")
        return ApiResponse.error(message=f"下载失败: {str(e)}")
    
    